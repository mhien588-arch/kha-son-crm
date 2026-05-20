"""
Tạo 2 file Hướng Dẫn — Kha Sơn Green Home
  1. Huong_Dan_Truong_Nhom.docx  — Hướng dẫn đầy đủ cho Hiển (trưởng nhóm)
  2. Huong_Dan_Nhan_Su.docx      — Hướng dẫn cơ bản cho nhân sự mới
Chạy: python make_huong_dan_v2.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── MÀU SẮC ─────────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x0D, 0x47, 0xA1)
NAVY_MID    = RGBColor(0x15, 0x65, 0xC0)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT   = RGBColor(0x21, 0x21, 0x21)
GREEN_OK    = RGBColor(0x2E, 0x7D, 0x32)
RED_WARN    = RGBColor(0xB7, 0x1C, 0x1C)
ORANGE      = RGBColor(0xE6, 0x5C, 0x00)
GOLD        = RGBColor(0xF5, 0x7F, 0x17)
GRAY_TEXT   = RGBColor(0x61, 0x61, 0x61)

H_NAVY      = "0D47A1"
H_NAVY_L    = "E3F2FD"
H_GREEN_L   = "E8F5E9"
H_ORANGE_L  = "FFF3E0"
H_RED_L     = "FFEBEE"
H_YELLOW_L  = "FFFDE7"
H_GRAY_L    = "F5F5F5"
H_WHITE     = "FFFFFF"


# ─── HELPERS ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_border(cell, color="0D47A1", sz="8"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ["top", "left", "bottom", "right"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=11,
            color=None, underline=False):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return r


def heading1(doc, text):
    """Tiêu đề cấp 1 — navy, lớn, bold, có đường kẻ dưới."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(4)
    r = add_run(p, text, bold=True, size=16, color=NAVY)
    p.paragraph_format.border_bottom = True
    # Kẻ dưới bằng XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement("w:bottom")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), "6")
    bdr.set(qn("w:color"), H_NAVY)
    pBdr.append(bdr)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    """Tiêu đề cấp 2 — navy mid, vừa, bold."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, bold=True, size=13, color=NAVY_MID)
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "▸ " + text, bold=True, size=11, color=DARK_TEXT)
    return p


def body(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=11, color=color or DARK_TEXT)
    return p


def bullet(doc, text, indent=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_run(p, text, size=10.5, color=color or DARK_TEXT)
    return p


def note_box(doc, text, bg=H_YELLOW_L, text_color=None):
    """Hộp ghi chú nền màu."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_border(cell, color="B0BEC5", sz="4")
    cell.width = Cm(15)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(4)
    cp.paragraph_format.left_indent = Cm(0.3)
    add_run(cp, text, size=10.5, color=text_color or DARK_TEXT)
    doc.add_paragraph()


def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement("w:bottom")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), "4")
    bdr.set(qn("w:color"), "90CAF9")
    pBdr.append(bdr)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def simple_table(doc, headers, rows, col_widths=None, header_bg=H_NAVY):
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"
    # Header
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        set_cell_border(cell, color=H_NAVY, sz="6")
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_before = Pt(3)
        cp.paragraph_format.space_after = Pt(3)
        add_run(cp, h, bold=True, size=10, color=WHITE)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row_data in enumerate(rows):
        drow = tbl.rows[ri + 1]
        bg = H_NAVY_L if ri % 2 == 0 else H_WHITE
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="90CAF9", sz="4")
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(2)
            if isinstance(val, tuple):
                text, bold, color = val
                add_run(cp, text, bold=bold, size=10, color=color)
            else:
                add_run(cp, str(val), size=10)
    if col_widths:
        for ri_all in range(len(tbl.rows)):
            for ci, w in enumerate(col_widths):
                tbl.rows[ri_all].cells[ci].width = Cm(w)
    doc.add_paragraph()


def page_break(doc):
    doc.add_page_break()


def set_margins(doc, top=2, bottom=2, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


# ═══════════════════════════════════════════════════════════════════════════════
# TÀI LIỆU 1 — HƯỚNG DẪN TRƯỞNG NHÓM (HIỂN)
# ═══════════════════════════════════════════════════════════════════════════════
def build_truong_nhom():
    doc = Document()
    set_margins(doc, top=2.2, bottom=2.2, left=2.8, right=2.5)

    # ── TRANG BÌA ─────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_title, "KHA SƠN GREEN HOME", bold=True, size=22, color=NAVY)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_sub, "HỆ THỐNG QUẢN LÝ KINH DOANH", bold=True, size=15, color=NAVY_MID)

    doc.add_paragraph()
    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_type, "HƯỚNG DẪN VẬN HÀNH TỔNG HỢP", bold=True, size=18, color=DARK_TEXT)
    add_run(p_type, "\nDÀNH CHO TRƯỞNG NHÓM", bold=True, size=13, color=GRAY_TEXT)

    for _ in range(2):
        doc.add_paragraph()

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_info, "Phiên bản: 17/05/2026  |  Người biên soạn: Hiển\n", size=10, color=GRAY_TEXT)
    add_run(p_info, "Dự án đất nền KCN Phú Bình, Thái Nguyên\n", size=10, color=GRAY_TEXT)
    add_run(p_info, "Hotline: 0368.557.832", size=10, color=GRAY_TEXT)

    page_break(doc)

    # ── 1. TỔNG QUAN HỆ THỐNG ─────────────────────────────────────────────────
    heading1(doc, "1. TỔNG QUAN HỆ THỐNG")
    body(doc, "Hệ thống CRM Kha Sơn Green Home được xây dựng để tự động hóa công việc quản lý khách hàng và theo dõi KPI cho team sale 6 người. Toàn bộ dữ liệu lưu trên Google Sheets, truy cập được từ bất kỳ thiết bị nào.")

    heading2(doc, "1.1 Cấu trúc nhân sự")
    simple_table(doc,
        ["#", "Tên", "Vai trò", "Nhiệm vụ chính"],
        [
            ("1", "Hiển", ("Trưởng nhóm", True, NAVY), "Dẫn dắt team, chốt deal lớn, báo cáo BGĐ"),
            ("2", "Đức", "Sale Senior", "Chăm sóc leads nóng, đưa khách đi xem đất"),
            ("3", "Sơn", "Sale Senior", "Chăm sóc leads nóng, đưa khách đi xem đất"),
            ("4", "Tuấn Anh", "Sale Junior", "Tìm kiếm leads mới, chăm sóc leads lạnh"),
            ("5", "Nam", "Sale Junior", "Tìm kiếm leads mới, hỗ trợ content FB/Zalo"),
            ("6", "Tiến", "Sale Junior", "Tìm kiếm leads mới, hỗ trợ content FB/Zalo"),
        ],
        col_widths=[1, 3, 4, 7]
    )

    heading2(doc, "1.2 KPI tháng")
    simple_table(doc,
        ["Chỉ tiêu", "Mục tiêu", "Ghi chú"],
        [
            ("Lead mới / tuần", "≥ 20 leads", "Đếm từ mọi nguồn: FB, Zalo, giới thiệu"),
            ("Lượt xem đất / tuần", "≥ 5 lượt", "Trạng thái 'Hẹn xem đất'"),
            ("Số cọc / tháng", "≥ 3 cọc", "Trạng thái 'Chốt cọc'"),
            ("CPL từ Ads", "< 100.000đ", "Chi phí Ads ÷ số lead từ Ads"),
        ],
        col_widths=[5, 4, 6]
    )

    heading2(doc, "1.3 Google Sheets CRM")
    body(doc, "Link: https://docs.google.com/spreadsheets/d/1syWcPK_YopGk2XtIYegDgHwnR6YPOHOoZmiN4ze-yjA")
    simple_table(doc,
        ["Tab", "Mục đích"],
        [
            ("KhaSonGreenHome_CRM", "Dữ liệu khách hàng chính — 54+ leads, định dạng navy"),
            ("SỐ MỚI KS", "Staging inbox — leads mới chờ sale liên hệ lần đầu + tích ✓"),
        ],
        col_widths=[6, 9]
    )

    heading2(doc, "1.4 Màu sắc trạng thái trên Google Sheets")
    simple_table(doc,
        ["Màu nền", "Trạng thái", "Ý nghĩa"],
        [
            ("Xanh dương nhạt", "Mới", "Lead vừa vào, chưa liên hệ"),
            ("Vàng nhạt", "Đang chăm sóc", "Đã liên hệ, đang theo dõi"),
            ("Cam nhạt", "Hẹn xem đất", "Đã hẹn lịch đi xem thực địa"),
            ("Xanh lá nhạt", "Chốt cọc", "Đã đặt cọc — THẮNG"),
            ("Xám nhạt", "Từ chối", "Không mua — đóng hồ sơ"),
        ],
        col_widths=[4, 4, 7]
    )

    divider(doc)

    # ── 2. CÁC FILE VẬN HÀNH ──────────────────────────────────────────────────
    heading1(doc, "2. CÁC FILE VẬN HÀNH")
    body(doc, "Tất cả file nằm trong thư mục D:\\QuanLyBKD2\\")
    body(doc, "Sale thông thường chỉ cần double-click file .bat. Trưởng nhóm có thể chạy Python trực tiếp để kiểm soát sâu hơn.")

    heading2(doc, "2.1 File .bat cho thao tác hằng ngày")
    simple_table(doc,
        ["File .bat", "Chạy script", "Dùng khi nào"],
        [
            ("Mo Tracker.bat", "tracker.py", "Xem + quản lý khách hàng — Dùng hằng ngày"),
            ("Bao Cao Ngay.bat", "daily_report.py", "Xem báo cáo cuối ngày — Chạy 18:00"),
            ("Bao Cao Tuan.bat", "weekly_report.py", "Báo cáo tuần — Chạy sáng thứ Hai"),
            ("Nhac Sang.bat", "nhac_sang.py", "Test nhắc nhở thủ công — Thay Task Scheduler tự chạy"),
            ("Don Dep Bao Cao.bat", "cleanup_reports.py", "Dọn dẹp file báo cáo cũ — Chạy đầu tháng"),
            ("Viet Content.bat", "content_helper.py", "Lấy kịch bản telesale / bài FB / Zalo"),
            ("Lich Ngay.bat", "lich_ngay.py", "Checklist công việc theo buổi sáng/chiều/tối"),
        ],
        col_widths=[4.5, 4.5, 6]
    )

    heading2(doc, "2.2 File cấu hình quan trọng")
    simple_table(doc,
        ["File", "Vai trò"],
        [
            ("config.json", "Nguồn cấu hình duy nhất: tên team, KPI, Sheet ID — KHÔNG xóa"),
            ("data/google_creds.json", "Khóa kết nối Google Sheets API — KHÔNG xóa, KHÔNG chia sẻ"),
        ],
        col_widths=[5, 10]
    )
    note_box(doc,
        "⚠  Lưu ý: Nếu mất file data/google_creds.json, hệ thống mất kết nối hoàn toàn với Google Sheets. "
        "Giữ bản backup trong ổ USB riêng.",
        bg=H_RED_L, text_color=RED_WARN
    )

    divider(doc)

    # ── 3. QUY TRÌNH VẬN HÀNH HÀNG NGÀY ─────────────────────────────────────
    heading1(doc, "3. QUY TRÌNH VẬN HÀNH HÀNG NGÀY")

    heading2(doc, "3.1 Buổi sáng (08:00 — 15 phút)")
    note_box(doc,
        "🔔  HỆ THỐNG TỰ ĐỘNG: Mỗi sáng 8:00 (T2–T7) Windows sẽ tự pop-up cửa sổ danh sách "
        "khách quá hạn. Nếu không có ai quá hạn, cửa sổ tự đóng.",
        bg=H_GREEN_L, text_color=GREEN_OK
    )
    bullet(doc, "Xem pop-up nhắc nhở tự động → ghi nhận danh sách khách cần gọi hôm nay")
    bullet(doc, "Double-click Mo Tracker.bat → nhấn [3] Kiểm tra khách bị quên → gọi ngay sale có tên trong cảnh báo")
    bullet(doc, "Kiểm tra Ads Manager: chi phí | số lead | CPL so với KPI")
    bullet(doc, "Nhấn [4] để thêm lead mới từ Facebook / Zalo / giới thiệu")
    bullet(doc, "Nhấn [6] để refresh dữ liệu và tự động chuyển lead đã tích ✓ từ 'SỐ MỚI KS' vào CRM chính")

    heading2(doc, "3.2 Buổi chiều (17:00 — 10 phút)")
    bullet(doc, "Double-click Mo Tracker.bat → nhấn [5] Cập nhật trạng thái lead sau khi sale gọi điện")
    bullet(doc, "Yêu cầu từng sale báo cáo kết quả cuộc gọi qua Zalo team")

    heading2(doc, "3.3 Cuối ngày (18:00 — 5 phút)")
    bullet(doc, "Double-click Bao Cao Ngay.bat → xem bảng KPI toàn team")
    bullet(doc, "Chụp màn hình → gửi group Zalo team")
    bullet(doc, "Lưu ý mục 'THẮT CỔ CHAI' trong phễu chuyển đổi để họp brief sáng hôm sau")

    heading2(doc, "3.4 Thứ Hai hằng tuần (08:30 — 10 phút)")
    bullet(doc, "Double-click Bao Cao Tuan.bat → báo cáo tuần (so sánh tuần này vs tuần trước)")
    bullet(doc, "Xem phân tích nguồn lead: Facebook / Zalo / Giới thiệu — nguồn nào tỷ lệ chốt cao nhất")
    bullet(doc, "Họp brief 15 phút với team: highlight lead nóng, phân công, mục tiêu tuần")

    heading2(doc, "3.5 Đầu tháng (ngày 1 hằng tháng — 5 phút)")
    bullet(doc, "Double-click Don Dep Bao Cao.bat → gõ 'yes' → file báo cáo tháng cũ được nén vào archive/")
    bullet(doc, "Xem leaderboard tháng vừa qua → thưởng / khích lệ team")

    divider(doc)

    # ── 4. HƯỚNG DẪN TRACKER.PY CHI TIẾT ───────────────────────────────────
    heading1(doc, "4. HƯỚNG DẪN CHI TIẾT TRACKER.PY")
    body(doc, "Chạy bằng: Mo Tracker.bat hoặc python tracker.py")

    heading2(doc, "Menu [1] — Xem tất cả khách hàng")
    bullet(doc, "Hiển thị bảng toàn bộ leads: màu đỏ = quá 3 ngày chưa liên hệ, xanh lá = Chốt cọc, xám = Từ chối")
    bullet(doc, "Số ngày ở cột cuối — sale ưu tiên gọi những lead số ngày cao nhất")

    heading2(doc, "Menu [2] — Lọc theo sale")
    bullet(doc, "Nhập tên sale (hoặc một phần tên) để xem toàn bộ leads của một người")
    bullet(doc, "Dùng khi muốn kiểm tra tiến độ từng cá nhân")

    heading2(doc, "Menu [3] — Kiểm tra khách bị quên (> 3 ngày)")
    bullet(doc, "Hiện đích danh tên sale + danh sách khách chưa được liên hệ quá 3 ngày")
    bullet(doc, "Đây là công cụ chính để nhắc nhở team mỗi sáng")

    heading2(doc, "Menu [4] — Thêm lead mới")
    bullet(doc, "Nhập: Tên khách → SĐT (bắt đầu 0, đủ 10 số) → Nguồn → Sale phụ trách")
    bullet(doc, "Hệ thống kiểm tra trùng SĐT cả trong CRM lẫn staging trước khi lưu")
    bullet(doc, "Lead vào tab 'SỐ MỚI KS' — sale cần tích ✓ sau khi liên hệ lần đầu")

    heading2(doc, "Menu [5] — Cập nhật trạng thái")
    bullet(doc, "Chọn số thứ tự khách → chọn trạng thái mới → cập nhật Nhu cầu nếu cần")
    bullet(doc, "Trạng thái hợp lệ: Mới | Đang chăm sóc | Hẹn xem đất | Chốt cọc | Từ chối")
    bullet(doc, "Cột 'Nhu cầu': Còn nhu cầu / Không còn nhu cầu — dùng để lọc leads 'nguội'")

    heading2(doc, "Menu [6] — Làm mới dữ liệu")
    bullet(doc, "Tải lại dữ liệu mới nhất từ Google Sheets")
    bullet(doc, "Tự động chuyển lead đã tích ✓ từ 'SỐ MỚI KS' vào CRM chính + in đậm dòng mới")

    heading2(doc, "Menu [7] — Quản lý team")
    bullet(doc, "[a] Thêm thành viên: nhập tên → tự cập nhật config.json")
    bullet(doc, "[c] Vô hiệu hóa: nhập tên → xác nhận 'yes' → xóa khỏi round-robin chia lead")
    bullet(doc, "Dữ liệu lịch sử của thành viên bị vô hiệu hóa vẫn giữ nguyên trong Google Sheets")

    heading2(doc, "Menu [8] — Tìm kiếm theo SĐT")
    bullet(doc, "Nhập số điện thoại (hoặc một phần) → hiện kết quả khớp")
    bullet(doc, "Nếu tìm thấy 1 khách → hỏi có muốn cập nhật trạng thái ngay không")

    heading2(doc, "Menu [9] — Xử lý SỐ MỚI KS")
    bullet(doc, "Chuyển thủ công các lead đã tích ✓ trong tab 'SỐ MỚI KS' vào CRM chính")
    bullet(doc, "Thường không cần dùng — hệ thống tự làm khi khởi động và nhấn [6]")

    divider(doc)

    # ── 5. LUỒNG XỬ LÝ LEAD MỚI ─────────────────────────────────────────────
    heading1(doc, "5. LUỒNG XỬ LÝ LEAD MỚI")
    simple_table(doc,
        ["Bước", "Ai làm", "Hành động", "Hệ thống ghi nhận"],
        [
            ("1", "Trưởng nhóm / Sale", "Nhập lead vào tracker menu [4]", "Lead vào tab 'SỐ MỚI KS', checkbox = FALSE"),
            ("2", "Sale phụ trách", "Gọi điện hoặc nhắn Zalo cho khách", "—"),
            ("3", "Sale phụ trách", "Mở Google Sheets → tích ✓ checkbox", "Checkbox = TRUE"),
            ("4", "Hệ thống tự động", "Khi tracker khởi động hoặc [6] refresh", "Lead chuyển sang CRM chính, in đậm dòng mới"),
            ("5", "Sale phụ trách", "Cập nhật trạng thái qua menu [5]", "Lưu ngày tương tác cuối"),
        ],
        col_widths=[1.5, 3.5, 5, 5]
    )
    note_box(doc,
        "💡  Khách chưa tích ✓ = sale chưa liên hệ lần đầu. "
        "Hệ thống sẽ KHÔNG đưa vào CRM cho đến khi sale xác nhận đã liên hệ.",
        bg=H_YELLOW_L
    )

    divider(doc)

    # ── 6. HỆ THỐNG BÁO CÁO ─────────────────────────────────────────────────
    heading1(doc, "6. HỆ THỐNG BÁO CÁO")

    heading2(doc, "6.1 Báo cáo ngày (Bao Cao Ngay.bat)")
    bullet(doc, "Block 1 — Tổng quan: Lead mới hôm nay | Đang chăm sóc | Chuyển đổi hôm nay")
    bullet(doc, "Block 2 — Từng sale: Đang giữ | Hẹn/Chốt hôm nay | Quá hạn (tên cụ thể)")
    bullet(doc, "Block 3 — Khách 'Không còn nhu cầu': cần xem lại hoặc đóng hồ sơ")
    bullet(doc, "Block 4 — Phễu chuyển đổi: thanh bar ASCII từ Mới → Chốt, highlight THẮT CỔ CHAI")
    bullet(doc, "Block 5 — Leaderboard tháng: điểm số từng sale (Chốt=10đ, Hẹn=3đ, 0 quá hạn=+2đ bonus)")
    bullet(doc, "Lưu tự động vào file: bao_cao_YYYY-MM-DD.txt")

    heading2(doc, "6.2 Báo cáo tuần (Bao Cao Tuan.bat)")
    bullet(doc, "Thống kê 7 ngày: Lead mới | Hẹn xem đất | Chốt cọc | Tỷ lệ chốt")
    bullet(doc, "So sánh tuần trước: +/- delta từng chỉ số")
    bullet(doc, "Sale hiệu quả nhất tuần")
    bullet(doc, "Phân tích nguồn: Facebook Ads / Zalo / Giới thiệu / Khác — tỷ lệ chốt từng nguồn")
    bullet(doc, "Lưu tự động vào: bao_cao_tuan_W{N}_2026.txt")

    heading2(doc, "6.3 Dọn dẹp báo cáo (Don Dep Bao Cao.bat)")
    bullet(doc, "Chạy đầu mỗi tháng — zip toàn bộ file báo cáo tháng cũ vào thư mục archive/")
    bullet(doc, "Preview danh sách file trước → xác nhận 'yes' → mới xóa và nén")
    bullet(doc, "File báo cáo tháng hiện tại KHÔNG bị xóa")

    divider(doc)

    # ── 7. HỆ THỐNG NHẮC NHỞ TỰ ĐỘNG ────────────────────────────────────────
    heading1(doc, "7. HỆ THỐNG NHẮC NHỞ TỰ ĐỘNG")

    heading2(doc, "7.1 Cách hoạt động")
    body(doc, "Windows Task Scheduler tự chạy script mỗi sáng 08:00 (T2–T7):")
    bullet(doc, "Nếu CÓ khách quá 3 ngày chưa liên hệ → pop-up cửa sổ đỏ, liệt kê theo từng sale, số ngày giảm dần")
    bullet(doc, "Nếu KHÔNG CÓ ai quá hạn → thoát luôn, không làm phiền")

    heading2(doc, "7.2 Kích hoạt lần đầu (chỉ làm 1 lần, cần quyền Admin)")
    bullet(doc, "Chuột phải vào file setup_task_scheduler.py → Run as administrator")
    bullet(doc, "Hoặc mở Command Prompt với quyền Admin → cd D:\\QuanLyBKD2 → python setup_task_scheduler.py")
    bullet(doc, "Kiểm tra: Windows + R → taskschd.msc → Task Scheduler Library → KhaSon_NhacNhoSang")

    heading2(doc, "7.3 Test thủ công")
    bullet(doc, "Double-click Nhac Sang.bat để xem kết quả ngay mà không cần đợi 08:00")

    note_box(doc,
        "⚡  Nếu máy tắt trước 08:00 → task không chạy hôm đó. "
        "Khuyến nghị: để máy tính bật từ 07:50 mỗi sáng.",
        bg=H_ORANGE_L, text_color=ORANGE
    )

    divider(doc)

    # ── 8. QUẢN LÝ TEAM ───────────────────────────────────────────────────────
    heading1(doc, "8. QUẢN LÝ TEAM SALE")

    heading2(doc, "8.1 Thêm thành viên mới")
    bullet(doc, "Mo Tracker.bat → nhấn [7] → nhấn [a] → nhập tên")
    bullet(doc, "Hệ thống tự cập nhật config.json — từ đây lead mới sẽ được chia đều cho cả người mới")

    heading2(doc, "8.2 Vô hiệu hóa thành viên")
    bullet(doc, "Mo Tracker.bat → nhấn [7] → nhấn [c] → chọn tên → gõ 'yes' để xác nhận")
    bullet(doc, "Dữ liệu khách hàng của người đó vẫn giữ nguyên trong Google Sheets")
    bullet(doc, "Cân nhắc chuyển khách hàng của họ sang sale khác trước khi vô hiệu hóa")

    heading2(doc, "8.3 Chia lead tự động")
    bullet(doc, "Hệ thống round-robin: chia đều cho từng sale theo thứ tự vòng tròn")
    bullet(doc, "Khi thêm sale mới, round-robin tự cập nhật → chia đều ngay lập tức")

    divider(doc)

    # ── 9. XỬ LÝ SỰ CỐ THƯỜNG GẶP ──────────────────────────────────────────
    heading1(doc, "9. XỬ LÝ SỰ CỐ THƯỜNG GẶP")
    simple_table(doc,
        ["Triệu chứng", "Nguyên nhân", "Cách xử lý"],
        [
            ("Script báo lỗi kết nối", "Mất Internet hoặc Google Sheets API timeout",
             "Kiểm tra Internet → chạy lại sau 30 giây"),
            ("SĐT bị mất số 0 đầu trong Sheet", "Sheets định dạng số học thay vì chuỗi",
             "Đã fix từ phiên 10 — nếu vẫn gặp: xóa ô + gõ lại, format cột = Văn bản"),
            ("Lead thêm vào không thấy trong CRM", "Lead đang ở tab 'SỐ MỚI KS' chờ tích ✓",
             "Mở Sheet → tab 'SỐ MỚI KS' → tích ✓ → chạy [6] hoặc khởi động lại tracker"),
            ("Task Scheduler không chạy lúc 08:00", "Máy tắt hoặc task bị disable",
             "Kiểm tra taskschd.msc → KhaSon_NhacNhoSang → Run Now để test"),
            ("Cột dữ liệu bị lệch trong tracker", "Ai đó xóa/thêm cột trong Google Sheets",
             "KHÔNG thêm/xóa cột trong Sheet — chỉ sửa dữ liệu trong các cột A-I có sẵn"),
            ("Script báo thiếu module", "Chưa cài thư viện Python",
             "Mở CMD: pip install pandas colorama gspread google-auth"),
        ],
        col_widths=[4.5, 4.5, 6]
    )

    divider(doc)

    # ── 10. QUY TẮC VÀNG ─────────────────────────────────────────────────────
    heading1(doc, "10. QUY TẮC VÀNG VẬN HÀNH")
    simple_table(doc,
        ["Tình huống", "Hành động bắt buộc"],
        [
            ("Lead mới < 1 giờ", "Gọi điện NGAY — không để qua ngày"),
            ("Khách chưa nghe máy", "Nhắn Zalo + gọi lại sau 2 tiếng"),
            ("Chưa liên hệ > 3 ngày", "Hệ thống tự nhắc — ưu tiên xử lý ngay buổi sáng"),
            ("Đã xem đất > 3 ngày chưa chốt", "Hỏi thăm, gửi thêm thông tin pháp lý/tiện ích"),
            ("Khách từ chối", "Ghi lý do vào Ghi chú → hỏi thăm lại sau 30 ngày"),
            ("Khách 'Không còn nhu cầu'", "Xem xét lại trong báo cáo ngày — có thể nurture dài hạn"),
            ("KHÔNG bao giờ", "Thêm / xóa cột trong Google Sheets hoặc xóa file google_creds.json"),
        ],
        col_widths=[5.5, 9.5]
    )

    page_break(doc)

    # ── PHỤ LỤC ───────────────────────────────────────────────────────────────
    heading1(doc, "PHỤ LỤC — CẤU TRÚC THƯ MỤC")
    body(doc, "D:\\QuanLyBKD2\\")
    simple_table(doc,
        ["File / Thư mục", "Mô tả"],
        [
            ("Mo Tracker.bat", "Launcher chính — double-click hằng ngày"),
            ("Bao Cao Ngay.bat", "Báo cáo cuối ngày"),
            ("Bao Cao Tuan.bat", "Báo cáo tuần"),
            ("Nhac Sang.bat", "Test nhắc nhở thủ công"),
            ("Don Dep Bao Cao.bat", "Dọn dẹp báo cáo tháng cũ"),
            ("Viet Content.bat", "Lấy kịch bản / bài đăng"),
            ("Lich Ngay.bat", "Checklist công việc"),
            ("config.json", "Cấu hình team + KPI + Sheet ID"),
            ("data/google_creds.json", "⚠ Credentials API — KHÔNG xóa"),
            ("archive/", "Báo cáo tháng cũ đã nén .zip"),
            ("bao_cao_*.txt", "Báo cáo ngày / tuần tháng hiện tại"),
            (".claude/commands/", "Claude Code skills (/lich-ngay, /bao-cao, ...)"),
        ],
        col_widths=[5, 10]
    )

    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# TÀI LIỆU 2 — HƯỚNG DẪN NHÂN SỰ MỚI
# ═══════════════════════════════════════════════════════════════════════════════
def build_nhan_su():
    doc = Document()
    set_margins(doc, top=2.5, bottom=2.5, left=3, right=2.5)

    # ── TRANG BÌA ─────────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p1, "KHA SƠN GREEN HOME", bold=True, size=20, color=NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p2, "HƯỚNG DẪN DÀNH CHO NHÂN SỰ", bold=True, size=17, color=DARK_TEXT)
    add_run(p2, "\n(Nhân viên mới và Sale)", size=12, color=GRAY_TEXT)

    for _ in range(3):
        doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p3, "Đất nền KCN Phú Bình, Thái Nguyên\n", size=11, color=GRAY_TEXT)
    add_run(p3, "Hotline: 0368.557.832  |  Cập nhật: 17/05/2026", size=10, color=GRAY_TEXT)

    page_break(doc)

    # ── 1. GIỚI THIỆU ─────────────────────────────────────────────────────────
    heading1(doc, "1. CHÀO MỪNG BẠN ĐẾN VỚI TEAM KHA SƠN!")
    body(doc, "Bạn vừa gia nhập team sale dự án đất nền Kha Sơn Green Home tại KCN Phú Bình, Thái Nguyên. "
         "Tài liệu này hướng dẫn bạn từng bước để bắt đầu làm việc hiệu quả ngay từ ngày đầu tiên.")

    heading2(doc, "Thông tin dự án")
    simple_table(doc,
        ["Hạng mục", "Chi tiết"],
        [
            ("Tên dự án", "Kha Sơn Green Home"),
            ("Vị trí", "Gần KCN Phú Bình, Thái Nguyên"),
            ("Sản phẩm", "Đất nền diện tích 80–120m²"),
            ("Giá từ", "800 triệu đồng"),
            ("Pháp lý", "Sổ đỏ lâu dài, minh bạch"),
            ("Hotline", "0368.557.832"),
        ],
        col_widths=[5, 10]
    )

    heading2(doc, "Đội ngũ của bạn")
    simple_table(doc,
        ["Tên", "Vai trò", "Liên hệ khi"],
        [
            ("Hiển", "Trưởng nhóm", "Mọi vấn đề kỹ thuật, quyết định quan trọng"),
            ("Đức", "Sale Senior", "Hỗ trợ chăm sóc khách hàng nóng"),
            ("Sơn", "Sale Senior", "Hỗ trợ chăm sóc khách hàng nóng"),
            ("Tuấn Anh", "Sale Junior", "Chia sẻ kinh nghiệm leads lạnh"),
            ("Nam", "Sale Junior", "Hỏi về content FB / Zalo"),
            ("Tiến", "Sale Junior", "Đồng nghiệp cùng cấp"),
        ],
        col_widths=[3.5, 4, 7.5]
    )

    divider(doc)

    # ── 2. CÁC TRẠNG THÁI KHÁCH HÀNG ─────────────────────────────────────────
    heading1(doc, "2. CÁC TRẠNG THÁI KHÁCH HÀNG — BẮT BUỘC NẮM VỮNG")
    body(doc, "Khi cập nhật thông tin khách hàng, bạn PHẢI chọn đúng một trong 5 trạng thái sau:")
    simple_table(doc,
        ["Trạng thái", "Nghĩa là", "Hành động tiếp theo"],
        [
            (("Mới", True, NAVY_MID), "Lead vừa vào, chưa ai liên hệ", "Gọi điện trong vòng 1 giờ"),
            (("Đang chăm sóc", True, GREEN_OK), "Đã liên hệ, đang theo dõi tiến độ", "Cập nhật ghi chú, hẹn gặp"),
            (("Hẹn xem đất", True, ORANGE), "Đã hẹn lịch đi thực địa", "Chuẩn bị thông tin, đưa đi xem"),
            (("Chốt cọc", True, GREEN_OK), "Khách đã đặt cọc — BẠN THẮNG!", "Làm hợp đồng, báo trưởng nhóm"),
            (("Từ chối", True, RED_WARN), "Khách không mua", "Ghi lý do, hẹn hỏi thăm sau 30 ngày"),
        ],
        col_widths=[4, 5, 6]
    )
    note_box(doc,
        "⚠  Không được để nguyên trạng thái 'Mới' quá 24 giờ. "
        "Lead nguội nhanh — gọi ngay trong ngày là tốt nhất!",
        bg=H_RED_L, text_color=RED_WARN
    )

    divider(doc)

    # ── 3. QUY TRÌNH HẰNG NGÀY ────────────────────────────────────────────────
    heading1(doc, "3. VIỆC BẠN CẦN LÀM MỖI NGÀY")

    heading2(doc, "Buổi sáng (08:00 — 15 phút)")
    bullet(doc, "Xem cửa sổ nhắc nhở tự động (nếu có) — đây là danh sách khách cần gọi hôm nay")
    bullet(doc, "Xem tin nhắn trong group Zalo team → cập nhật tình hình")
    bullet(doc, "Mở Google Sheets (link trưởng nhóm sẽ gửi) → kiểm tra tab 'SỐ MỚI KS' → có lead mới giao cho bạn không?")

    heading2(doc, "Trong ngày — Gọi điện & nhắn tin")
    bullet(doc, "Ưu tiên gọi lead mới nhất trước — lead nóng nguội rất nhanh")
    bullet(doc, "Sau mỗi cuộc gọi: cập nhật trạng thái ngay trong Google Sheets hoặc báo trưởng nhóm")
    bullet(doc, "Nếu khách hứa gặp → đổi trạng thái sang 'Hẹn xem đất' + ghi ngày hẹn vào Ghi chú")

    heading2(doc, "Buổi chiều (17:00 — 5 phút)")
    bullet(doc, "Báo cáo kết quả gọi điện hôm nay qua group Zalo: bao nhiêu cuộc, ai hứa gặp, ai từ chối")
    bullet(doc, "Trưởng nhóm hoặc bạn cập nhật trạng thái trong hệ thống")

    divider(doc)

    # ── 4. CÁCH SỬ DỤNG GOOGLE SHEETS ────────────────────────────────────────
    heading1(doc, "4. CÁCH SỬ DỤNG GOOGLE SHEETS")
    body(doc, "Đây là nguồn dữ liệu duy nhất của team. Bạn có thể mở trên điện thoại hoặc máy tính bất kỳ.")

    heading2(doc, "4.1 Truy cập")
    bullet(doc, "Trưởng nhóm sẽ chia sẻ link Google Sheets với bạn qua Zalo")
    bullet(doc, "Đăng nhập bằng tài khoản Google của bạn → Yêu cầu quyền chỉnh sửa nếu cần")
    bullet(doc, "Bookmark link để truy cập nhanh")

    heading2(doc, "4.2 Tab 'KhaSonGreenHome_CRM' — Dữ liệu chính")
    bullet(doc, "Cột A: Tên khách hàng")
    bullet(doc, "Cột B: Ngày tiếp cận lần đầu")
    bullet(doc, "Cột C: Số điện thoại — KHÔNG xóa số 0 đầu!")
    bullet(doc, "Cột D: Nguồn (Facebook Ads / Zalo / Giới thiệu / Khác)")
    bullet(doc, "Cột E: Trạng thái (xem Phần 2 ở trên)")
    bullet(doc, "Cột F: Ghi chú — ghi đặc điểm khách, lý do từ chối, hẹn lịch cụ thể")
    bullet(doc, "Cột G: Sale chăm sóc — tên bạn")
    bullet(doc, "Cột H: Ngày tương tác cuối — hệ thống dùng để tính ngày quá hạn")
    bullet(doc, "Cột I: Nhu cầu — 'Còn nhu cầu' hoặc 'Không còn nhu cầu'")

    heading2(doc, "4.3 Tab 'SỐ MỚI KS' — Leads mới chờ xử lý")
    note_box(doc,
        "Khi bạn thấy tên mình trong tab 'SỐ MỚI KS' — đó là lead mới được giao. "
        "Hãy gọi điện NGAY rồi tích ✓ vào ô checkbox cột I.",
        bg=H_GREEN_L, text_color=GREEN_OK
    )
    bullet(doc, "Cột I (cuối cùng): Checkbox — tích ✓ khi bạn đã liên hệ lần đầu")
    bullet(doc, "Sau khi tích ✓, hệ thống sẽ tự động chuyển lead sang CRM chính trong lần refresh tiếp theo")
    bullet(doc, "CHƯA liên hệ = CHƯA tích ✓ — tuyệt đối không tích trước khi gọi")

    heading2(doc, "4.4 Cách cập nhật dữ liệu trên điện thoại")
    bullet(doc, "Mở link Google Sheets → chạm vào ô Trạng thái (cột E) → chọn trạng thái mới từ dropdown")
    bullet(doc, "Cập nhật cột Ghi chú (F) và Ngày tương tác cuối (H) sau mỗi cuộc gọi")
    bullet(doc, "Tích ✓ checkbox trong tab 'SỐ MỚI KS' sau khi liên hệ lần đầu")

    divider(doc)

    # ── 5. QUY TẮC CHĂM SÓC KHÁCH HÀNG ──────────────────────────────────────
    heading1(doc, "5. QUY TẮC CHĂM SÓC KHÁCH HÀNG")
    simple_table(doc,
        ["Tình huống", "Hành động"],
        [
            ("Lead mới (< 1 giờ)", "Gọi điện NGAY, không để qua ngày — khách nhớ bạn nhất lúc vừa để lại thông tin"),
            ("Khách chưa nghe máy", "Nhắn Zalo + gọi lại sau 2 tiếng"),
            ("Khách hứa gọi lại", "Đặt reminder trong điện thoại, gọi đúng giờ đã hẹn"),
            ("Chưa liên hệ > 3 ngày", "Hệ thống sẽ báo tên bạn — xử lý ngay trong ngày"),
            ("Đã xem đất > 3 ngày", "Hỏi thăm, gửi thêm thông tin: pháp lý, tiện ích xung quanh"),
            ("Khách từ chối", "KHÔNG tranh cãi. Ghi lý do vào Ghi chú → hỏi thăm sau 30 ngày"),
            ("Khách hỏi giá", "Giá từ 800 triệu, diện tích 80–120m² — Mời xem thực địa trước"),
        ],
        col_widths=[5, 10]
    )

    divider(doc)

    # ── 6. KPI CÁ NHÂN ────────────────────────────────────────────────────────
    heading1(doc, "6. KPI CÁ NHÂN CỦA BẠN")
    body(doc, "Mỗi sale cần đóng góp vào KPI chung của team:")
    simple_table(doc,
        ["Chỉ tiêu", "Mục tiêu team", "Phần bạn cần đóng góp"],
        [
            ("Lead mới / tuần", "≥ 20 leads", "≥ 3–4 leads / tuần"),
            ("Lượt xem đất / tuần", "≥ 5 lượt", "≥ 1 lượt / tuần"),
            ("Số cọc / tháng", "≥ 3 cọc", "≥ 0–1 cọc / tháng (tùy cấp bậc)"),
        ],
        col_widths=[5, 4, 6]
    )

    heading2(doc, "Hệ thống tính điểm Leaderboard")
    simple_table(doc,
        ["Hành động", "Điểm thưởng"],
        [
            ("Chốt cọc thành công (mỗi lead)", "+10 điểm"),
            ("Dẫn khách đi xem đất (mỗi lượt)", "+3 điểm"),
            ("0 khách quá hạn trong tháng", "+2 điểm bonus"),
        ],
        col_widths=[9, 6]
    )
    body(doc, "Xem leaderboard: trong báo cáo ngày (Bao Cao Ngay.bat) — Block 5 cuối báo cáo.")

    divider(doc)

    # ── 7. KỊCH BẢN TELESALE ─────────────────────────────────────────────────
    heading1(doc, "7. KỊCH BẢN GỌI ĐIỆN — GỢI Ý")
    note_box(doc,
        "💡  Đây chỉ là gợi ý. Hãy nói theo phong cách tự nhiên của bạn — khách cảm nhận được sự chân thành hơn kịch bản đọc vẹt.",
        bg=H_NAVY_L
    )

    heading2(doc, "Kịch bản A — Lead từ Facebook Ads")
    body(doc, '"Alo, em chào anh/chị [Tên]. Em là [Tên sale] từ dự án Kha Sơn Green Home. '
         'Hôm qua anh/chị có để lại thông tin quan tâm đất nền gần KCN Phú Bình đúng không ạ? '
         'Dạ, dự án mình đang có lô đẹp từ 800 triệu, sổ đỏ lâu dài. '
         'Không biết anh/chị đang tìm đất để đầu tư hay để xây nhà ở ạ?"')

    heading2(doc, "Kịch bản B — Gọi lại sau khi khách đã xem thông tin")
    body(doc, '"Alo anh/chị [Tên] ơi, em [Tên sale] đây ạ. '
         'Hôm trước anh/chị có xem thông tin Kha Sơn Green Home — '
         'dạo này anh/chị đã có thêm quyết định gì chưa ạ? '
         'Em vừa có lô góc đẹp vừa ra, muốn báo ngay để anh/chị cân nhắc sớm."')

    heading2(doc, "Kịch bản C — Mời xem thực địa")
    body(doc, '"Anh/chị [Tên] ơi, để anh/chị có cái nhìn thực tế nhất, '
         'em có thể sắp xếp đưa anh/chị ra xem đất cuối tuần này được không? '
         'Chỉ cần 30–45 phút thôi ạ, ra tận nơi là anh/chị sẽ thấy tiềm năng rõ hơn nhiều."')

    divider(doc)

    # ── 8. NHỮNG ĐIỀU QUAN TRỌNG CẦN NHỚ ────────────────────────────────────
    heading1(doc, "8. NHỮNG ĐIỀU QUAN TRỌNG PHẢI NHỚ")

    note_box(doc, "✅  CẦN LÀM:\n"
             "• Cập nhật trạng thái trong Google Sheets sau MỖI cuộc gọi\n"
             "• Tích ✓ checkbox khi đã liên hệ lần đầu lead mới\n"
             "• Báo cáo kết quả trong group Zalo team cuối ngày\n"
             "• Hỏi trưởng nhóm nếu không chắc về bước tiếp theo",
             bg=H_GREEN_L, text_color=GREEN_OK)

    note_box(doc, "❌  KHÔNG được làm:\n"
             "• Thêm / xóa cột trong Google Sheets (làm hỏng hệ thống)\n"
             "• Tích ✓ checkbox trước khi thực sự gọi điện\n"
             "• Để lead > 3 ngày không liên hệ mà không báo trưởng nhóm\n"
             "• Chia sẻ link file data/google_creds.json cho người ngoài",
             bg=H_RED_L, text_color=RED_WARN)

    divider(doc)

    # ── 9. LIÊN HỆ HỖ TRỢ ────────────────────────────────────────────────────
    heading1(doc, "9. KHI BẠN CẦN HỖ TRỢ")
    simple_table(doc,
        ["Loại vấn đề", "Liên hệ"],
        [
            ("Hỏi về sản phẩm / dự án / pháp lý", "Hiển (Trưởng nhóm) — Hotline: 0368.557.832"),
            ("Hệ thống báo lỗi / không chạy được", "Hiển (Trưởng nhóm) hoặc gửi ảnh chụp màn hình vào group"),
            ("Không thấy lead trong Google Sheets", "Kiểm tra tab 'SỐ MỚI KS' trước → báo Hiển nếu vẫn không thấy"),
            ("Muốn xem kịch bản telesale / bài đăng FB", "Double-click Viet Content.bat trên máy tính của Hiển"),
            ("Khách phức tạp, không biết xử lý", "Báo ngay Hiển hoặc Đức / Sơn (Sale Senior)"),
        ],
        col_widths=[6, 9]
    )

    for _ in range(2):
        doc.add_paragraph()

    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_final, "Chúc bạn bán hàng thành công! 💪", bold=True, size=14, color=NAVY)
    add_run(p_final, "\nKha Sơn Green Home — KCN Phú Bình, Thái Nguyên", size=10, color=GRAY_TEXT)

    return doc


# ─── MAIN ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))
    out1 = os.path.join(base, "Huong_Dan_Truong_Nhom.docx")
    out2 = os.path.join(base, "Huong_Dan_Nhan_Su.docx")

    print("  Dang tao Huong_Dan_Truong_Nhom.docx ...")
    doc1 = build_truong_nhom()
    doc1.save(out1)
    print(f"  [OK] {out1}")

    print("  Dang tao Huong_Dan_Nhan_Su.docx ...")
    doc2 = build_nhan_su()
    doc2.save(out2)
    print(f"  [OK] {out2}")

    print("\n  Hoan tat! 2 file da duoc tao.")
